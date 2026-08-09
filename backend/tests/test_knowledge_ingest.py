"""ADR-017 stage 17.1: backend/knowledge_ingest.py - the extract -> chunk ->
store pipeline, and its extension-dispatch onto backend/attachments.py's
own extraction functions."""

from __future__ import annotations

import pytest

from backend.attachments import AttachmentError
from backend.knowledge_ingest import IngestError, extract_text, ingest_file, ingest_text


@pytest.fixture
def db_path(tmp_path):
    return tmp_path / "knowledge.db"


# -- extract_text: extension dispatch ----------------------------------------


class TestExtractTextDispatch:
    def test_plain_text_extension_reads_raw_text(self, tmp_path):
        path = tmp_path / "notes.txt"
        path.write_text("Plain content here.")
        text, mime = extract_text(path)
        assert text == "Plain content here."
        assert mime == "text/plain"

    def test_markdown_extension_reads_as_plain_text(self, tmp_path):
        # ADR-017's own extraction-reuse note: .md is already in
        # attachments.PLAIN_TEXT_EXTENSIONS - no structural markdown
        # parsing in this stage, just the same raw-text read every other
        # PLAIN_TEXT_EXTENSIONS member gets. write_bytes (not write_text):
        # _read_text decodes raw bytes with no newline translation, and
        # Path.write_text would silently widen \n to \r\n on Windows.
        path = tmp_path / "readme.md"
        path.write_bytes(b"# Heading\n\nBody text.")
        text, mime = extract_text(path)
        assert text == "# Heading\n\nBody text."
        assert mime == "text/plain"

    def test_code_extension_reads_as_plain_text(self, tmp_path):
        path = tmp_path / "script.py"
        path.write_text("def f():\n    return 1\n")
        text, mime = extract_text(path)
        assert "def f():" in text
        assert mime == "text/plain"

    def test_csv_extension_reads_as_plain_text(self, tmp_path):
        path = tmp_path / "data.csv"
        path.write_bytes(b"a,b,c\n1,2,3\n")
        text, mime = extract_text(path)
        assert text == "a,b,c\n1,2,3\n"
        assert mime == "text/plain"

    def test_docx_extension_reuses_attachments_pys_real_extraction(self, tmp_path):
        pytest.importorskip("docx")
        import docx

        path = tmp_path / "doc.docx"
        document = docx.Document()
        document.add_paragraph("First real paragraph.")
        document.add_paragraph("Second real paragraph.")
        document.save(str(path))

        text, mime = extract_text(path)
        assert "First real paragraph." in text
        assert "Second real paragraph." in text
        assert mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def test_pdf_with_no_extractable_text_is_rejected_not_silently_empty(self, tmp_path):
        pytest.importorskip("pypdf")
        from pypdf import PdfWriter

        path = tmp_path / "blank.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=72, height=72)
        with open(path, "wb") as f:
            writer.write(f)

        with pytest.raises(AttachmentError, match="No readable text could be extracted"):
            extract_text(path)

    def test_missing_file_raises_ingest_error(self, tmp_path):
        with pytest.raises(IngestError, match="File not found"):
            extract_text(tmp_path / "does-not-exist.txt")

    def test_unsupported_extension_raises_ingest_error(self, tmp_path):
        path = tmp_path / "image.png"
        path.write_bytes(b"\x89PNG\r\n\x1a\n" + b"\x00" * 20)
        with pytest.raises(IngestError, match="Unsupported file type"):
            extract_text(path)


# -- HTML extraction: new in this stage, not a pass-through of raw tags -----


class TestHtmlExtraction:
    def test_html_strips_nav_and_script_and_keeps_semantic_content(self, tmp_path):
        pytest.importorskip("bs4")
        path = tmp_path / "page.html"
        path.write_text(
            "<html><head><title>T</title><script>evil()</script></head>"
            "<body><nav>Skip this nav</nav>"
            "<main><h1>Real Heading</h1><p>Real paragraph content.</p></main>"
            "<footer>Skip this footer</footer></body></html>"
        )
        text, mime = extract_text(path)
        assert mime == "text/html"
        assert "Real Heading" in text
        assert "Real paragraph content." in text
        assert "Skip this nav" not in text
        assert "Skip this footer" not in text
        assert "evil()" not in text

    def test_html_with_no_semantic_tags_falls_back_to_whole_document_text(self, tmp_path):
        pytest.importorskip("bs4")
        path = tmp_path / "bare.html"
        path.write_text("<html><body><div>Just a bare div, no semantic tags at all.</div></body></html>")
        text, mime = extract_text(path)
        assert "Just a bare div, no semantic tags at all." in text

    def test_html_with_no_readable_text_after_stripping_raises(self, tmp_path):
        pytest.importorskip("bs4")
        path = tmp_path / "empty.html"
        path.write_text("<html><head><script>only script content</script></head><body></body></html>")
        with pytest.raises(IngestError, match="no readable text"):
            extract_text(path)


# -- ingest_file: the full extract -> chunk -> store pipeline ---------------


class TestIngestFile:
    def test_ingesting_a_text_file_produces_a_document_with_chunks(self, db_path, tmp_path):
        path = tmp_path / "notes.txt"
        path.write_text("First paragraph of real content.\n\nSecond paragraph with more words.")
        outcome = ingest_file(str(path), db_path=db_path)
        assert outcome.was_new is True
        assert outcome.chunk_count >= 1

        from backend.knowledge_store import get_document, list_chunks_for_document
        doc = get_document(db_path, outcome.document_id)
        assert doc["title"] == "notes.txt"
        assert doc["mime"] == "text/plain"
        chunks = list_chunks_for_document(db_path, outcome.document_id)
        assert len(chunks) == outcome.chunk_count

    def test_reingesting_the_same_file_is_a_no_op(self, db_path, tmp_path):
        path = tmp_path / "notes.txt"
        path.write_text("Identical content every time.")
        first = ingest_file(str(path), db_path=db_path)
        second = ingest_file(str(path), db_path=db_path)
        assert second.was_new is False
        assert second.document_id == first.document_id

    def test_ingesting_into_different_collections_creates_separate_documents(self, db_path, tmp_path):
        path = tmp_path / "notes.txt"
        path.write_text("Shared file content.")
        first = ingest_file(str(path), db_path=db_path, collection_id=1)
        second = ingest_file(str(path), db_path=db_path, collection_id=2)
        assert first.document_id != second.document_id

    def test_a_file_with_no_indexable_text_raises_ingest_error(self, db_path, tmp_path):
        # Empty file: readable (an empty string IS valid plain text), but
        # chunk_text() returns no chunks for it - must surface as a clear
        # ingest-time error, not a silently-empty document with zero chunks.
        path = tmp_path / "empty.txt"
        path.write_text("")
        with pytest.raises(IngestError, match="no indexable text"):
            ingest_file(str(path), db_path=db_path)

    def test_custom_chunking_parameters_are_honored(self, db_path, tmp_path):
        path = tmp_path / "long.txt"
        paragraphs = [f"Paragraph {i} with several filler words for bulk." for i in range(20)]
        path.write_text("\n\n".join(paragraphs))

        default_outcome = ingest_file(str(path), db_path=db_path, collection_id=1)
        small_target_outcome = ingest_file(
            str(path), db_path=db_path, collection_id=2, target_tokens=20, overlap_tokens=5,
        )
        # A tighter token budget must produce strictly more chunks for the
        # SAME source text.
        assert small_target_outcome.chunk_count > default_outcome.chunk_count


# -- ingest_text: web-research retention / branch indexing's own entry point


class TestIngestText:
    def test_ingesting_text_produces_a_document_with_chunks(self, db_path):
        outcome = ingest_text(
            "First paragraph of real content.\n\nSecond paragraph with more words.",
            source_uri="https://example.com/page", title="Example Page", db_path=db_path,
        )
        assert outcome.was_new is True
        assert outcome.chunk_count >= 1

        from backend.knowledge_store import list_chunks_for_document
        chunks = list_chunks_for_document(db_path, outcome.document_id)
        assert len(chunks) == outcome.chunk_count

    def test_ingesting_text_stores_the_given_source_uri_title_and_mime(self, db_path):
        outcome = ingest_text(
            "Some retained web content here.",
            source_uri="https://example.com/article", title="Article Title",
            mime="text/html", db_path=db_path,
        )
        from backend.knowledge_store import get_document
        doc = get_document(db_path, outcome.document_id)
        assert doc["source_uri"] == "https://example.com/article"
        assert doc["title"] == "Article Title"
        assert doc["mime"] == "text/html"

    def test_reingesting_the_same_text_is_a_no_op(self, db_path):
        first = ingest_text("Identical retained content.", source_uri="u1", title="t1", db_path=db_path)
        second = ingest_text("Identical retained content.", source_uri="u2", title="t2", db_path=db_path)
        # Idempotency is by content hash (+ collection), matching
        # ingest_file()'s own contract - a DIFFERENT source_uri/title for
        # the SAME text is still the same document.
        assert second.was_new is False
        assert second.document_id == first.document_id

    def test_blank_text_raises_ingest_error(self, db_path):
        with pytest.raises(IngestError, match="no indexable text"):
            ingest_text("   ", source_uri="u", title="Empty", db_path=db_path)

    def test_default_mime_is_text_plain(self, db_path):
        outcome = ingest_text("Plain retained text.", source_uri="u", title="t", db_path=db_path)
        from backend.knowledge_store import get_document
        assert get_document(db_path, outcome.document_id)["mime"] == "text/plain"
