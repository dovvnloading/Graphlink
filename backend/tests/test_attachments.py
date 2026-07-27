"""Real coverage for backend/attachments.py (R8a).

This module is a faithful port of the deleted graphlink_app/
graphlink_file_handler.py + graphlink_window.py's _stage_attachment_file
(git show 6c919f6~1:...) - these tests exercise the SAME classification
decisions against real files on disk, not mocked ones, since the whole
point of this feature was that the previous "it works" claim (a DOM node
existing) was never actually tested against real behavior.
"""

from __future__ import annotations

import os

import pytest

from backend.attachments import AttachmentError, stage_file


@pytest.fixture
def tmp_file(tmp_path):
    def _write(name: str, content: bytes | str) -> str:
        path = tmp_path / name
        if isinstance(content, str):
            path.write_text(content, encoding="utf-8")
        else:
            path.write_bytes(content)
        return str(path)

    return _write


class TestImageClassification:
    def test_png_is_classified_as_image_with_real_bytes_in_the_content_part(self, tmp_file):
        real_png_header = b"\x89PNG\r\n\x1a\n" + b"\x00" * 32
        path = tmp_file("photo.png", real_png_header)

        staged = stage_file(path)

        assert staged.kind == "image"
        assert staged.context_label == "Vision"
        assert staged.content_part == {"type": "image_bytes", "data": real_png_header}
        assert staged.extracted_text is None

    @pytest.mark.parametrize("ext", [".png", ".jpg", ".jpeg", ".webp"])
    def test_every_documented_image_extension_classifies_as_image(self, tmp_file, ext):
        path = tmp_file(f"pic{ext}", b"\x00" * 16)
        assert stage_file(path).kind == "image"


class TestAudioClassification:
    def test_a_readable_wav_is_classified_as_audio_with_a_path_content_part(self, tmp_file):
        # A minimal but real, playable WAV: RIFF/WAVE header + one silent
        # PCM sample, so graphlink_audio.inspect_audio_file's real duration
        # probe succeeds rather than needing a mock.
        import struct
        import tempfile
        import wave

        fd, real_path = tempfile.mkstemp(suffix=".wav")
        os.close(fd)
        with wave.open(real_path, "wb") as w:
            w.setnchannels(1)
            w.setsampwidth(2)
            w.setframerate(8000)
            w.writeframes(struct.pack("<h", 0) * 8000)  # 1 second of silence

        try:
            staged = stage_file(real_path)
            assert staged.kind == "audio"
            assert staged.content_part == {"type": "audio_file", "path": real_path}
            assert staged.duration_seconds == pytest.approx(1.0, abs=0.05)
            assert "Audio" in staged.context_label
            assert staged.extracted_text is None
        finally:
            os.unlink(real_path)

    def test_audio_over_the_duration_cap_is_rejected_not_silently_truncated(self, tmp_file, monkeypatch):
        import backend.attachments as attachments_module

        def _fake_inspect(path):
            raise attachments_module.AudioValidationError(
                "'huge.wav' is 5:00:00 long. The maximum supported length is 4:00:00."
            )

        monkeypatch.setattr(attachments_module, "inspect_audio_file", _fake_inspect)
        path = tmp_file("huge.wav", b"RIFF....WAVEfmt ")

        with pytest.raises(AttachmentError, match="maximum supported length"):
            stage_file(path)


class TestDocumentClassificationAndExtraction:
    def test_plain_text_extension_is_read_verbatim_with_a_real_token_count(self, tmp_file):
        path = tmp_file("notes.txt", "Hello, this is a real attachment body.")

        staged = stage_file(path)

        assert staged.kind == "document"
        assert staged.context_label == "Text"
        assert staged.extracted_text == "Hello, this is a real attachment body."
        assert staged.token_count > 0
        assert staged.content_part is None

    def test_code_extension_gets_the_code_label(self, tmp_file):
        path = tmp_file("script.py", "def f():\n    return 1\n")
        assert stage_file(path).context_label == "Code"

    def test_markdown_extension_gets_the_markdown_label(self, tmp_file):
        path = tmp_file("readme.md", "# Title\n\nBody text.")
        assert stage_file(path).context_label == "Markdown"

    def test_unrecognized_extension_with_text_like_bytes_is_sniffed_and_read(self, tmp_file):
        # No extension in PLAIN_TEXT_EXTENSIONS, but the byte-sniff heuristic
        # (ported verbatim from FileHandler._looks_like_text_file) should
        # still classify and read it as a document.
        path = tmp_file("data.xyz", "just plain ascii text content for sniffing purposes")

        staged = stage_file(path)

        assert staged.kind == "document"
        assert staged.extracted_text is not None

    def test_named_file_with_no_extension_is_recognized(self, tmp_file):
        path = tmp_file("Dockerfile", "FROM python:3.12\n")
        staged = stage_file(path)
        assert staged.kind == "document"

    def test_binary_garbage_with_unknown_extension_is_rejected(self, tmp_file):
        path = tmp_file("blob.xyz", bytes(range(256)) * 20)

        with pytest.raises(AttachmentError, match="Unsupported file type"):
            stage_file(path)

    def test_null_byte_containing_file_is_rejected_even_with_mostly_ascii_content(self, tmp_file):
        # The byte-sniff heuristic's own hard rule: any NUL byte disqualifies
        # a file outright, matching FileHandler._looks_like_text_file exactly.
        path = tmp_file("mixed.bin", b"mostly text here\x00but one null byte")
        with pytest.raises(AttachmentError, match="Unsupported file type"):
            stage_file(path)


class TestPdfAndDocx:
    def test_missing_pypdf_surfaces_the_exact_install_message(self, tmp_file, monkeypatch):
        import builtins

        real_import = builtins.__import__

        def _blocked_import(name, *args, **kwargs):
            if name in ("pypdf", "PyPDF2"):
                raise ImportError(name)
            return real_import(name, *args, **kwargs)

        monkeypatch.setattr(builtins, "__import__", _blocked_import)
        path = tmp_file("doc.pdf", b"%PDF-1.4 fake")

        with pytest.raises(AttachmentError, match="pip install pypdf"):
            stage_file(path)

    def test_missing_python_docx_surfaces_the_exact_install_message(self, tmp_file, monkeypatch):
        import builtins

        real_import = builtins.__import__

        def _blocked_import(name, *args, **kwargs):
            if name == "docx":
                raise ImportError(name)
            return real_import(name, *args, **kwargs)

        monkeypatch.setattr(builtins, "__import__", _blocked_import)
        path = tmp_file("doc.docx", b"PK\x03\x04 fake docx bytes")

        with pytest.raises(AttachmentError, match="pip install python-docx"):
            stage_file(path)

    def test_a_real_docx_extracts_its_real_paragraph_text(self, tmp_file):
        pytest.importorskip("docx")
        import docx

        import tempfile

        fd, real_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        document = docx.Document()
        document.add_paragraph("First real paragraph.")
        document.add_paragraph("Second real paragraph.")
        document.save(real_path)

        try:
            staged = stage_file(real_path)
            assert staged.kind == "document"
            assert staged.context_label == "DOCX"
            assert "First real paragraph." in staged.extracted_text
            assert "Second real paragraph." in staged.extracted_text
        finally:
            os.unlink(real_path)

    def test_a_real_pdf_with_no_extractable_text_is_rejected_not_silently_empty(self, tmp_file):
        pytest.importorskip("pypdf")
        from pypdf import PdfWriter

        import tempfile

        fd, real_path = tempfile.mkstemp(suffix=".pdf")
        os.close(fd)
        writer = PdfWriter()
        writer.add_blank_page(width=72, height=72)  # a real page, zero text
        with open(real_path, "wb") as f:
            writer.write(f)

        try:
            with pytest.raises(AttachmentError, match="No readable text could be extracted"):
                stage_file(real_path)
        finally:
            os.unlink(real_path)


class TestRejections:
    def test_missing_file_is_rejected(self, tmp_path):
        with pytest.raises(AttachmentError, match="File not found"):
            stage_file(str(tmp_path / "does-not-exist.txt"))

    def test_a_directory_is_rejected_not_read_as_a_file(self, tmp_path):
        with pytest.raises(AttachmentError, match="File not found"):
            stage_file(str(tmp_path))
